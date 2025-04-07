import os
import logging
import tempfile
import io
import requests
import re
from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup, ReplyKeyboardMarkup, BotCommand
from telegram.ext import Application, CommandHandler, MessageHandler, CallbackQueryHandler, ContextTypes, filters
from dotenv import load_dotenv
from PIL import Image
from pdf2docx import Converter
from docx import Document
from PyPDF2 import PdfReader, PdfWriter

load_dotenv()
BOT_TOKEN = os.getenv("BOT_TOKEN")
CLOUD_CONVERT_API_KEY = os.getenv("CLOUD_CONVERT_API_KEY", "")
ADMIN_ID = os.getenv("ADMIN_ID", "145414784")

logging.basicConfig(
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s", level=logging.INFO
)
logger = logging.getLogger(__name__)

def get_main_keyboard():
    keyboard = [
        ["🔄 Fayllarni o'zgartirish"],
        ["📄 Betlash"],
        ["🔤 Almashtirish"]
    ]
    return ReplyKeyboardMarkup(keyboard, resize_keyboard=True)

def get_convert_keyboard():
    keyboard = [
        [
            InlineKeyboardButton("PDF → Word", callback_data="pdf_to_word"),
            InlineKeyboardButton("Word → PDF", callback_data="word_to_pdf")
        ],
        [InlineKeyboardButton("⬅️ Orqaga", callback_data="back_to_main")]
    ]
    return InlineKeyboardMarkup(keyboard)

def get_transliteration_keyboard():
    keyboard = [
        [
            InlineKeyboardButton("Kril → Lotin", callback_data="cyrillic_to_latin"),
            InlineKeyboardButton("Lotin → Kril", callback_data="latin_to_cyrillic")
        ],
        [InlineKeyboardButton("⬅️ Orqaga", callback_data="back_to_main")]
    ]
    return InlineKeyboardMarkup(keyboard)

async def start(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    user = update.effective_user
    
    if not context.user_data.get('notified_admin'):
        user_info = f"Yangi foydalanuvchi botni ishga tushirdi!\n\n"
        user_info += f"ID: {user.id}\n"
        user_info += f"Username: @{user.username if user.username else 'username yo\'q'}\n"
        user_info += f"Ism: {user.first_name} {user.last_name if user.last_name else ''}"
        
        try:
            await context.bot.send_message(chat_id=ADMIN_ID, text=user_info)
            context.user_data['notified_admin'] = True
        except Exception as e:
            logger.error(f"Error sending message to admin: {str(e)}")
    
    await update.message.reply_html(
        f"Salom, {user.mention_html()}! 👋\n\n"
        f"Men PDF va Word formatlarini o'zgartirishda yordam beraman.\n\n"
        f"Quyidagi variantlardan birini tanlang:",
        reply_markup=get_main_keyboard(),
    )

async def help_command(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    await update.message.reply_text(
        "Men quyidagi ishlarni bajarishda yordam beraman:\n\n"
        "1. PDF fayllarni Word hujjatlariga o'zgartirish\n"
        "2. Word hujjatlarini PDF fayllariga o'zgartirish\n"
        "3. PDF yoki Word fayllardan belgilangan betlarni ajratib olish\n"
        "4. Fayllarni Kril va Lotin alifbosida almashtirish\n\n"
        "Quyidagi menyudan variantni tanlang:",
        reply_markup=get_main_keyboard()
    )

async def menu_command(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    await update.message.reply_text(
        "Asosiy menyu:\n\n"
        "Quyidagi variantlardan birini tanlang:",
        reply_markup=get_main_keyboard()
    )

async def handle_text(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    text = update.message.text
    
    if text == "🔄 Fayllarni o'zgartirish":
        await update.message.reply_text(
            "O'zgartirish turini tanlang:",
            reply_markup=get_convert_keyboard()
        )
    elif text == "📄 Betlash":
        context.user_data['waiting_for_file'] = 'page_selection'
        await update.message.reply_text(
            "Iltimos, betlarni ajratib olmoqchi bo'lgan PDF yoki Word faylni yuklang."
        )
    elif text == "🔤 Almashtirish":
        await update.message.reply_text(
            "Almashtirish turini tanlang:",
            reply_markup=get_transliteration_keyboard()
        )
    elif context.user_data.get('waiting_for_pages'):
        await handle_page_input(update, context)
    else:
        await update.message.reply_text(
            "Iltimos, menyudan variantni tanlang:",
            reply_markup=get_main_keyboard()
        )

async def handle_page_input(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    page_input = update.message.text
    file_path = context.user_data.get('file_path')
    file_type = context.user_data.get('file_type')
    
    if not file_path or not os.path.exists(file_path):
        await update.message.reply_text(
            "Fayl topilmadi. Iltimos, faylni qaytadan yuboring.",
            reply_markup=get_main_keyboard()
        )
        context.user_data['waiting_for_pages'] = False
        return
    
    try:
        pages_to_extract = parse_page_ranges(page_input)
        
        if not pages_to_extract:
            await update.message.reply_text(
                "Noto'g'ri format. Misol uchun: 1-5, 7, 10-12",
                reply_markup=get_main_keyboard()
            )
            return
        
        await update.message.reply_text(f"Quyidagi betlarni ajratyapman: {', '.join(map(str, pages_to_extract))}")
        
        if file_type == 'pdf':
            output_bytes = await extract_pdf_pages(file_path, pages_to_extract)
            output_name = os.path.basename(file_path).replace('.pdf', '_selected_pages.pdf')
            await update.message.reply_document(
                document=output_bytes,
                filename=output_name,
                caption=f"Tanlangan betlar: {', '.join(map(str, pages_to_extract))}"
            )
        elif file_type == 'docx':
            output_bytes = await extract_docx_pages(file_path, pages_to_extract)
            output_name = os.path.basename(file_path).replace('.docx', '_selected_pages.docx')
            await update.message.reply_document(
                document=output_bytes,
                filename=output_name,
                caption=f"Tanlangan betlar: {', '.join(map(str, pages_to_extract))}"
            )
        
    except Exception as e:
        logger.error(f"Error processing page selection: {str(e)}")
        await update.message.reply_text(
            f"Kechirasiz, betlarni ajratishda xatolik yuz berdi: {str(e)}",
            reply_markup=get_main_keyboard()
        )
    finally:
        if os.path.exists(file_path):
            os.unlink(file_path)
        context.user_data['waiting_for_pages'] = False
        context.user_data['file_path'] = None
        context.user_data['file_type'] = None
        
        await update.message.reply_text(
            "Betlarni ajratish tugallandi! Yana nima qilmoqchisiz?",
            reply_markup=get_main_keyboard()
        )

def parse_page_ranges(page_input):
    pages = []
    ranges = page_input.replace(' ', '').split(',')
    
    for r in ranges:
        if '-' in r:
            try:
                start, end = map(int, r.split('-'))
                if start <= end:
                    pages.extend(range(start, end + 1))
            except ValueError:
                continue
        else:
            try:
                pages.append(int(r))
            except ValueError:
                continue
    
    return sorted(set(pages))

async def extract_pdf_pages(pdf_path, pages_to_extract):
    output_path = pdf_path.replace('.pdf', '_selected_pages.pdf')
    
    try:
        pdf_reader = PdfReader(pdf_path)
        pdf_writer = PdfWriter()
        
        total_pages = len(pdf_reader.pages)
        valid_pages = [p for p in pages_to_extract if 1 <= p <= total_pages]
        
        if not valid_pages:
            raise Exception(f"Tanlangan betlar mavjud emas. Fayl {total_pages} betdan iborat.")
        
        for page_num in valid_pages:
            pdf_writer.add_page(pdf_reader.pages[page_num - 1])
        
        with open(output_path, 'wb') as output_file:
            pdf_writer.write(output_file)
        
        with open(output_path, 'rb') as file:
            output_bytes = file.read()
        
        return io.BytesIO(output_bytes)
    
    except Exception as e:
        logger.error(f"Error extracting PDF pages: {str(e)}")
        raise
    finally:
        if os.path.exists(output_path):
            os.unlink(output_path)

async def extract_docx_pages(docx_path, pages_to_extract):
    pdf_path = docx_path.replace('.docx', '_temp.pdf')
    output_pdf_path = docx_path.replace('.docx', '_selected_pages.pdf')
    output_docx_path = docx_path.replace('.docx', '_selected_pages.docx')
    
    try:
        doc = Document(docx_path)
        
        if CLOUD_CONVERT_API_KEY:
            with open(docx_path, 'rb') as docx_file:
                pdf_bytes = await convert_word_to_pdf_cloud(docx_file.read(), os.path.basename(docx_path))
            
            with open(pdf_path, 'wb') as pdf_file:
                pdf_file.write(pdf_bytes.getbuffer())
        else:
            from fpdf import FPDF
            pdf = FPDF()
            
            for para in doc.paragraphs:
                pdf.add_page()
                pdf.set_font("Arial", size=12)
                text = para.text.encode('latin-1', 'replace').decode('latin-1')
                pdf.multi_cell(0, 10, text)
            
            pdf.output(pdf_path)
        
        pdf_reader = PdfReader(pdf_path)
        pdf_writer = PdfWriter()
        
        total_pages = len(pdf_reader.pages)
        valid_pages = [p for p in pages_to_extract if 1 <= p <= total_pages]
        
        if not valid_pages:
            raise Exception(f"Tanlangan betlar mavjud emas. Fayl {total_pages} betdan iborat.")
        
        for page_num in valid_pages:
            pdf_writer.add_page(pdf_reader.pages[page_num - 1])
        
        with open(output_pdf_path, 'wb') as output_file:
            pdf_writer.write(output_file)
        
        cv = Converter(output_pdf_path)
        cv.convert(output_docx_path, start=0, end=None)
        cv.close()
        
        with open(output_docx_path, 'rb') as file:
            output_bytes = file.read()
        
        return io.BytesIO(output_bytes)
    
    except Exception as e:
        logger.error(f"Error extracting DOCX pages: {str(e)}")
        raise
    finally:
        for path in [pdf_path, output_pdf_path, output_docx_path]:
            if os.path.exists(path):
                os.unlink(path)

def cyrillic_to_latin(text):
    cyrillic_to_latin_map = {
        'а': 'a', 'б': 'b', 'в': 'v', 'г': 'g', 'д': 'd', 'е': 'e', 'ё': 'yo',
        'ж': 'j', 'з': 'z', 'и': 'i', 'й': 'y', 'к': 'k', 'л': 'l', 'м': 'm',
        'н': 'n', 'о': 'o', 'п': 'p', 'р': 'r', 'с': 's', 'т': 't', 'у': 'u',
        'ф': 'f', 'х': 'x', 'ц': 'ts', 'ч': 'ch', 'ш': 'sh', 'щ': 'sh', 'ъ': ''',
        'ы': 'i', 'ь': '', 'э': 'e', 'ю': 'yu', 'я': 'ya', 'ў': 'o'', 'қ': 'q',
        'ғ': 'g'', 'ҳ': 'h',
        'А': 'A', 'Б': 'B', 'В': 'V', 'Г': 'G', 'Д': 'D', 'Е': 'E', 'Ё': 'Yo',
        'Ж': 'J', 'З': 'Z', 'И': 'I', 'Й': 'Y', 'К': 'K', 'Л': 'L', 'М': 'M',
        'Н': 'N', 'О': 'O', 'П': 'P', 'Р': 'R', 'С': 'S', 'Т': 'T', 'У': 'U',
        'Ф': 'F', 'Х': 'X', 'Ц': 'Ts', 'Ч': 'Ch', 'Ш': 'Sh', 'Щ': 'Sh', 'Ъ': ''',
        'Ы': 'I', 'Ь': '', 'Э': 'E', 'Ю': 'Yu', 'Я': 'Ya', 'Ў': 'O\'', 'Қ': 'Q',
        'Ғ': 'G\'', 'Ҳ': 'H'
    }
    
    text = text.replace('Ц', 'S')
    text = text.replace('ц', 's')
    
    result = ""
    for char in text:
        result += cyrillic_to_latin_map.get(char, char)
    
    return result

def latin_to_cyrillic(text):
    text = re.sub(r"O['']", "Ў", text)
    text = re.sub(r"o['']", "ў", text)
    text = re.sub(r"G['']", "Ғ", text)
    text = re.sub(r"g['']", "ғ", text)
    text = re.sub(r"Ch", "Ч", text)
    text = re.sub(r"ch", "ч", text)
    text = re.sub(r"Sh", "Ш", text)
    text = re.sub(r"sh", "ш", text)
    text = re.sub(r"Yu", "Ю", text)
    text = re.sub(r"yu", "ю", text)
    text = re.sub(r"Ya", "Я", text)
    text = re.sub(r"ya", "я", text)
    text = re.sub(r"Yo", "Ё", text)
    text = re.sub(r"yo", "ё", text)
    text = re.sub(r"Ts", "Ц", text)
    text = re.sub(r"ts", "ц", text)
    
    latin_to_cyrillic_map = {
        'a': 'а', 'b': 'б', 'v': 'в', 'g': 'г', 'd': 'д', 'e': 'е',
        'j': 'ж', 'z': 'з', 'i': 'и', 'y': 'й', 'k': 'к', 'l': 'л', 'm': 'м',
        'n': 'н', 'o': 'о', 'p': 'п', 'r': 'р', 's': 'с', 't': 'т', 'u': 'у',
        'f': 'ф', 'x': 'х', 'h': 'ҳ', 'q': 'қ',
        'A': 'А', 'B': 'Б', 'V': 'В', 'G': 'Г', 'D': 'Д', 'E': 'Е',
        'J': 'Ж', 'Z': 'З', 'I': 'И', 'Y': 'Й', 'K': 'К', 'L': 'Л', 'M': 'М',
        'N': 'Н', 'O': 'О', 'P': 'П', 'R': 'Р', 'S': 'С', 'T': 'Т', 'U': 'У',
        'F': 'Ф', 'X': 'Х', 'H': 'Ҳ', 'Q': 'Қ'
    }
    
    result = ""
    i = 0
    while i < len(text):
        char = text[i]
        result += latin_to_cyrillic_map.get(char, char)
        i += 1
    
    return result

async def transliterate_docx(file_bytes, file_name, to_latin=True):
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_docx:
        temp_docx.write(file_bytes)
        docx_path = temp_docx.name
    
    suffix = "_to_latin" if to_latin else "_to_cyrillic"
    output_path = docx_path.replace('.docx', f'{suffix}.docx')
    
    try:
        doc = Document(docx_path)
        
        for para in doc.paragraphs:
            if to_latin:
                para.text = cyrillic_to_latin(para.text)
            else:
                para.text = latin_to_cyrillic(para.text)
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if to_latin:
                            paragraph.text = cyrillic_to_latin(paragraph.text)
                        else:
                            paragraph.text = latin_to_cyrillic(paragraph.text)
        
        doc.save(output_path)
        
        with open(output_path, 'rb') as file:
            output_bytes = file.read()
        
        return io.BytesIO(output_bytes)
    
    except Exception as e:
        logger.error(f"Error transliterating DOCX: {str(e)}")
        raise
    finally:
        if os.path.exists(docx_path):
            os.unlink(docx_path)
        if os.path.exists(output_path):
            os.unlink(output_path)

async def transliterate_pdf(file_bytes, file_name, to_latin=True):
    with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_pdf:
        temp_pdf.write(file_bytes)
        pdf_path = temp_pdf.name
    
    docx_path = pdf_path.replace('.pdf', '_temp.docx')
    suffix = "_to_latin" if to_latin else "_to_cyrillic"
    output_docx_path = pdf_path.replace('.pdf', f'{suffix}.docx')
    output_pdf_path = pdf_path.replace('.pdf', f'{suffix}.pdf')
    
    try:
        cv = Converter(pdf_path)
        cv.convert(docx_path)
        cv.close()
        
        doc = Document(docx_path)
        
        for para in doc.paragraphs:
            if to_latin:
                para.text = cyrillic_to_latin(para.text)
            else:
                para.text = latin_to_cyrillic(para.text)
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if to_latin:
                            paragraph.text = cyrillic_to_latin(paragraph.text)
                        else:
                            paragraph.text = latin_to_cyrillic(paragraph.text)
        
        doc.save(output_docx_path)
        
        if file_name.lower().endswith('.pdf'):
            if CLOUD_CONVERT_API_KEY:
                with open(output_docx_path, 'rb') as docx_file:
                    pdf_bytes = await convert_word_to_pdf_cloud(docx_file.read(), os.path.basename(output_docx_path))
                
                with open(output_pdf_path, 'wb') as pdf_file:
                    pdf_file.write(pdf_bytes.getbuffer())
                
                with open(output_pdf_path, 'rb') as file:
                    output_bytes = file.read()
                
                return io.BytesIO(output_bytes)
            else:
                from fpdf import FPDF
                pdf = FPDF()
                
                for para in doc.paragraphs:
                    pdf.add_page()
                    pdf.set_font("Arial", size=12)
                    text = para.text.encode('latin-1', 'replace').decode('latin-1')
                    pdf.multi_cell(0, 10, text)
                
                pdf.output(output_pdf_path)
                
                with open(output_pdf_path, 'rb') as file:
                    output_bytes = file.read()
                
                return io.BytesIO(output_bytes)
        else:
            with open(output_docx_path, 'rb') as file:
                output_bytes = file.read()
            
            return io.BytesIO(output_bytes)
    
    except Exception as e:
        logger.error(f"Error transliterating PDF: {str(e)}")
        raise
    finally:
        for path in [pdf_path, docx_path, output_docx_path, output_pdf_path]:
            if os.path.exists(path):
                os.unlink(path)

async def handle_document(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    file = update.message.document
    file_name = file.file_name
    
    if not context.user_data.get('waiting_for_file'):
        await update.message.reply_text(
            "Iltimos, avval asosiy menyudan nima qilmoqchi ekanligingizni tanlang.",
            reply_markup=get_main_keyboard()
        )
        return
    
    file_type = context.user_data.get('waiting_for_file')
    
    if file_type == 'pdf_to_word' and not file_name.lower().endswith('.pdf'):
        await update.message.reply_text(
            "Iltimos, PDF faylni yuklang.",
            reply_markup=get_convert_keyboard()
        )
        return
    
    if file_type == 'word_to_pdf' and not file_name.lower().endswith(('.docx', '.doc')):
        await update.message.reply_text(
            "Iltimos, Word hujjatini (DOCX/DOC) yuklang.",
            reply_markup=get_convert_keyboard()
        )
        return
    
    if file_type in ['page_selection', 'cyrillic_to_latin', 'latin_to_cyrillic'] and not file_name.lower().endswith(('.pdf', '.docx', '.doc')):
        await update.message.reply_text(
            "Iltimos, PDF yoki Word (DOCX/DOC) faylni yuklang.",
            reply_markup=get_main_keyboard()
        )
        return
    
    await update.message.reply_text("Faylingiz qayta ishlanmoqda, iltimos kuting...")
    
    try:
        new_file = await context.bot.get_file(file.file_id)
        file_bytes = await new_file.download_as_bytearray()
        
        if file_type == 'pdf_to_word':
            output_bytes = await convert_pdf_to_word(file_bytes, file_name)
            output_name = file_name.replace('.pdf', '.docx')
            await update.message.reply_document(
                document=output_bytes,
                filename=output_name,
                caption="Mana sizning Word hujjatingiz!"
            )
            context.user_data['waiting_for_file'] = None
            
        elif file_type == 'word_to_pdf':
            output_bytes = await convert_word_to_pdf_cloud(file_bytes, file_name)
            output_name = file_name.replace('.docx', '').replace('.doc', '') + '.pdf'
            await update.message.reply_document(
                document=output_bytes,
                filename=output_name,
                caption="Mana sizning PDF faylingiz!"
            )
            context.user_data['waiting_for_file'] = None
            
        elif file_type == 'page_selection':
            with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(file_name)[1]) as temp_file:
                temp_file.write(file_bytes)
                file_path = temp_file.name
            
            context.user_data['file_path'] = file_path
            context.user_data['file_type'] = 'pdf' if file_name.lower().endswith('.pdf') else 'docx'
            context.user_data['waiting_for_pages'] = True
            context.user_data['waiting_for_file'] = None
            
            await update.message.reply_text(
                "Betlarni kiriting. Misol uchun: 1-5, 7, 10-12\n\n"
                "Qaysi betlarni ajratib olishni istaysiz?"
            )
            return
            
        elif file_type == 'cyrillic_to_latin':
            if file_name.lower().endswith('.docx'):
                output_bytes = await transliterate_docx(file_bytes, file_name, to_latin=True)
                output_name = file_name.replace('.docx', '_to_latin.docx')
            else:
                output_bytes = await transliterate_pdf(file_bytes, file_name, to_latin=True)
                output_name = file_name.replace('.pdf', '_to_latin.pdf')
                
            await update.message.reply_document(
                document=output_bytes,
                filename=output_name,
                caption="Mana sizning Lotincha faylingiz!"
            )
            context.user_data['waiting_for_file'] = None
            
        elif file_type == 'latin_to_cyrillic':
            if file_name.lower().endswith('.docx'):
                output_bytes = await transliterate_docx(file_bytes, file_name, to_latin=False)
                output_name = file_name.replace('.docx', '_to_cyrillic.docx')
            else:
                output_bytes = await transliterate_pdf(file_bytes, file_name, to_latin=False)
                output_name = file_name.replace('.pdf', '_to_cyrillic.pdf')
                
            await update.message.reply_document(
                document=output_bytes,
                filename=output_name,
                caption="Mana sizning Kirilcha faylingiz!"
            )
            context.user_data['waiting_for_file'] = None
        
        await update.message.reply_text(
            "O'zgartirish tugallandi! Yana nima qilmoqchisiz?",
            reply_markup=get_main_keyboard()
        )
    except Exception as e:
        logger.error(f"Error processing file: {str(e)}")
        await update.message.reply_text(
            f"Kechirasiz, faylni qayta ishlashda xatolik yuz berdi: {str(e)}",
            reply_markup=get_main_keyboard()
        )
        context.user_data['waiting_for_file'] = None

async def convert_pdf_to_word(file_bytes, file_name):
    with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_pdf:
        temp_pdf.write(file_bytes)
        pdf_path = temp_pdf.name
    
    docx_path = pdf_path.replace('.pdf', '.docx')
    
    try:
        cv = Converter(pdf_path)
        cv.convert(docx_path)
        cv.close()
        
        with open(docx_path, 'rb') as f:
            output_bytes = f.read()
        
        return io.BytesIO(output_bytes)
    except Exception as e:
        logger.error(f"Error converting PDF to Word: {str(e)}")
        raise
    finally:
        if os.path.exists(pdf_path):
            os.unlink(pdf_path)
        if os.path.exists(docx_path):
            os.unlink(docx_path)

async def convert_word_to_pdf_cloud(file_bytes, file_name):
    if not CLOUD_CONVERT_API_KEY:
        raise Exception(
            "Cloud Convert API kalit so'zi sozlanmagan. Iltimos, .env fayliga CLOUD_CONVERT_API_KEY qo'shing. "
            "https://cloudconvert.com/ saytidan bepul kalit oling."
        )
    
    try:
        url = "https://api.cloudconvert.com/v2/jobs"
        payload = {
            "tasks": {
                "upload-my-file": {
                    "operation": "import/upload"
                },
                "convert-my-file": {
                    "operation": "convert",
                    "input": "upload-my-file",
                    "output_format": "pdf",
                    "some_other_option": "value"
                },
                "export-my-file": {
                    "operation": "export/url",
                    "input": "convert-my-file"
                }
            }
        }
        
        headers = {
            "Authorization": f"Bearer {CLOUD_CONVERT_API_KEY}",
            "Content-Type": "application/json"
        }
        
        response = requests.post(url, json=payload, headers=headers)
        response.raise_for_status()
        data = response.json()
        
        upload_url = data["data"]["tasks"][0]["result"]["form"]["url"]
        upload_params = data["data"]["tasks"][0]["result"]["form"]["parameters"]
        
        files = {
            "file": (file_name, io.BytesIO(file_bytes), "application/octet-stream")
        }
        
        response = requests.post(upload_url, data=upload_params, files=files)
        response.raise_for_status()
        
        job_id = data["data"]["id"]
        wait_url = f"https://api.cloudconvert.com/v2/jobs/{job_id}/wait"
        
        response = requests.get(wait_url, headers=headers)
        response.raise_for_status()
        data = response.json()
        
        export_task = None
        for task in data["data"]["tasks"]:
            if task["name"] == "export-my-file" and task["status"] == "finished":
                export_task = task
                break
        
        if not export_task or "result" not in export_task or "files" not in export_task["result"]:
            raise Exception("Konvertatsiya jarayonida xatolik yuz berdi.")
        
        download_url = export_task["result"]["files"][0]["url"]
        
        response = requests.get(download_url)
        response.raise_for_status()
        
        return io.BytesIO(response.content)
    
    except requests.exceptions.RequestException as e:
        logger.error(f"Error in CloudConvert API: {str(e)}")
        if hasattr(e, 'response') and e.response is not None:
            logger.error(f"Response: {e.response.text}")
        raise Exception(f"Cloud konvertatsiya xizmatida xatolik: {str(e)}")
    except Exception as e:
        logger.error(f"Error converting Word to PDF with cloud service: {str(e)}")
        raise

async def handle_callback(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    query = update.callback_query
    await query.answer()
    
    if query.data == "back_to_main":
        await query.message.reply_text(
            "Nima qilmoqchisiz?",
            reply_markup=get_main_keyboard()
        )
        if 'waiting_for_file' in context.user_data:
            del context.user_data['waiting_for_file']
    
    elif query.data == "pdf_to_word":
        context.user_data['waiting_for_file'] = 'pdf_to_word'
        await query.message.reply_text(
            "Iltimos, Word formatiga o'zgartirmoqchi bo'lgan PDF faylni yuklang."
        )
    
    elif query.data == "word_to_pdf":
        context.user_data['waiting_for_file'] = 'word_to_pdf'
        await query.message.reply_text(
            "Iltimos, PDF formatiga o'zgartirmoqchi bo'lgan Word hujjatini (DOCX/DOC) yuklang."
        )
    
    elif query.data == "cyrillic_to_latin":
        context.user_data['waiting_for_file'] = 'cyrillic_to_latin'
        await query.message.reply_text(
            "Iltimos, Kirildan Lotinga o'zgartirmoqchi bo'lgan PDF yoki Word faylni yuklang."
        )
    
    elif query.data == "latin_to_cyrillic":
        context.user_data['waiting_for_file'] = 'latin_to_cyrillic'
        await query.message.reply_text(
            "Iltimos, Lotindan Kirilga o'zgartirmoqchi bo'lgan PDF yoki Word faylni yuklang."
        )

async def setup_commands(application):
    commands = [
        BotCommand("start", "Botni ishga tushirish"),
        BotCommand("help", "Yordam olish"),
    ]
    
    await application.bot.set_my_commands(commands)

def main() -> None:
    application = Application.builder().token(BOT_TOKEN).build()

    application.add_handler(CommandHandler("start", start))
    application.add_handler(CommandHandler("help", help_command))
    application.add_handler(CommandHandler("menu", menu_command))
    application.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, handle_text))
    application.add_handler(MessageHandler(filters.Document.ALL, handle_document))
    application.add_handler(CallbackQueryHandler(handle_callback))

    application.post_init = setup_commands

    print("Bot started...")
    application.run_polling()

if __name__ == "__main__":
    main()